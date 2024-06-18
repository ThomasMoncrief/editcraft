import docx, json, re
from django.contrib.auth import authenticate, login, logout
from django.shortcuts import render
from django.http import HttpResponse, HttpResponseRedirect, JsonResponse
from django.urls import reverse
from django.views.decorators.csrf import csrf_exempt
from copyeditor.functions import run_editor, get_title, compare_text, create_html
# from django.views.decorators.cache import cache_control

from .models import User, Archive

def index(request):
    if not request.user.is_authenticated:
        return render(request, "login.html", {
                "message": "You must be logged in.",
                "return_to": "index"
            })

    else:
        return render(request, "index.html")


def login_view(request):
    if request.method == "POST":
        username = request.POST["username"]
        password = request.POST["password"]
        user = authenticate(request, username=username, password=password)
        return_to = request.POST["return_to"]

        if user is not None:
            login(request, user)
            try:
                return HttpResponseRedirect(reverse(return_to))
            except:
                return HttpResponseRedirect(reverse("index"))
        else:
            return render(request, "login.html", {
                "message": "Invalid username and/or password."
            })
    else:
        return render(request, "login.html")


def logout_view(request):
    logout(request)
    return HttpResponseRedirect(reverse("index"))


def signup(request):
    if request.method == "POST":
        username = request.POST["username"]
        email = request.POST["email"]
        password = request.POST["password"]
        confirmation = request.POST["confirmation"]
        if password != confirmation:
            return render(request, "signup.html", {
                "alert": "Passwords do not match."
            })

        user = User.objects.create_user(username, email, password)
        user.save()
        return HttpResponseRedirect(reverse("index"))
    else:
        return render(request, "signup.html")


def settings(request):
    if not request.user.is_authenticated:
        return render(request, "login.html", {
                "message": "Log in to view your settings.",
                "return_to": "settings"
            })
    
    user = User.objects.get(username=request.user)
    
    if request.method == "POST":
        user.key = request.POST["api-key"]
        user.save()
        return HttpResponseRedirect(reverse("settings"))

    return render(request, "settings.html", {
        "user": request.user,
        "key": user.key
    })


def uploader(request):
    if not request.user.is_authenticated:
        return render(request, "login.html", {
                "message": "Log in to begin uploading",
                "return_to": "uploader"
            })
    if request.method == "POST":
        # clear submit_text in case back button clicked
        submit_text = ""
        
        # Handle uploaded file
        if request.POST.get('action') == 'upload_file':
            try:
                file = request.FILES['upload_file']
            except:
                return HttpResponse("Must upload a file")
            
            extension = file.name.split(".")[1]
            if extension == "txt":
                print(".txt file uploaded")
                submit_text = file.read().decode('utf-8', errors='ignore')
            elif extension == "docx":
                print(".docx file type uploaded")
                file = docx.Document(file)
                for paragraph in file.paragraphs:
                    submit_text += paragraph.text + "\n"
            else:
                return HttpResponse("Invalid file type.")

        # Handle pasted text
        elif request.POST["text_box"]:
            submit_text = request.POST["text_box"]
            submit_text = re.sub(r'\r', '', submit_text) #eliminate carriage returns submitted through HMTL

        # Submit the text to the LLM
        key = User.objects.get(username=request.user).key
        result = run_editor(submit_text, key)
        
        # Process results
        if result == "key invalid":
            return HttpResponse("Invalid OpenAI key. Make sure to set your key in Settings.")
        if result[0] == "":
            result = result[1:] #eliminate first line if it is blank
        title = get_title(result)
        diffs = compare_text(submit_text, result)
        save_in_archive = Archive(user=request.user, title=title, original_text=submit_text, edited_text=result, diffs=diffs)
        save_in_archive.save()

        return HttpResponseRedirect(f"workshop/{save_in_archive.id}")

    return render(request, "uploader.html")


@csrf_exempt
def workshop(request):
    #generic workshop view
    if not request.user.is_authenticated:
        return render(request, "login.html", {
                "message": "You must be logged in to view the Workshop.",
                "return_to": "workshop"
            })
    
    articles = Archive.objects.filter(user=request.user.id)
    return render(request, "workshop.html", {
        "articles": articles
    })


def workshop_render(request, id):
    #This method will need to be changed later so that users cannot edit each
    # other's posts
    if not request.user.is_authenticated:
        return render(request, "login.html", {
                "message": "You must be logged in to view the Workshop.",
                "return_to": "workshop"
            })
    archive = Archive.objects.get(id=id)
    
    edited_text = archive.edited_text
    edited_text = re.sub(r'\r', '', edited_text)
    
    # compare_text() will always compare unaccepted changes in the final text against the GPT edit.
    unedited_text = archive.final_text
    if unedited_text == "":
        unedited_text = archive.original_text #if no final text exists yet, grab the original text
    
    unedited_text = re.sub(r'\r', '', unedited_text)
    
    #**
    diffs = archive.diffs

    # if archive.user == request.user:
    preview_text = create_html(diffs)
    return render(request, "workshop_render.html", {
        "text": preview_text,
        "article_id": id,
        "title": archive.title
        })
    # else: **WILL TURN THIS BACK ON WHEN TESTING IS FINISHED 
    # return HttpResponse("You do not have access to that post")

@csrf_exempt
def workshop_download(request, id):

    if not request.user.is_authenticated:
        return render(request, "login.html", {
                "message": "You must be logged in to view the Workshop.",
                "return_to": "workshop"
            })
    final_text_string = Archive.objects.get(id=id).final_text
    file_type = request.GET.get('type')
    
    if file_type == "txt":
        response = HttpResponse(final_text_string, content_type='text/plain')
        response['Content-Disposition'] = 'attachment; filename="edited_text.txt"'
        return response
    
    if file_type == "docx":
        edited = docx.Document()
        for paragraph in final_text_string.split("\n"):
            edited.add_paragraph(paragraph)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename = "edited_text.docx"'
        edited.save(response)
        return response
    
    return HttpResponseRedirect(reverse("workshop"))


@csrf_exempt
# @cache_control(no_cache=True, must_revalidate=True, no_store=True)
def workshop_api(request, id):
    
    # delete article, sent from workshop.js
    if request.method == "DELETE":
        Archive.objects.get(id=id).delete()
        return HttpResponseRedirect(reverse("workshop"))
    
    # sent from workshop_render.js
    elif request.method == "PUT":
        article = Archive.objects.get(id=id)
        data = json.loads(request.body)
        
        if data["request_type"] == "update_title":
            article.title = data["new_title"]

        elif data["request_type"] == "update_text":
            final_text = data["final_text"].strip()
            final_text = re.sub(r'\r', '', final_text) #eliminate carriage returns submitted through HMTL
            article.final_text = final_text
            article.diffs = compare_text(final_text, article.edited_text)

        elif data["request_type"] == "revert_changes":
            article.final_text = "" #if no final text exists, original text is used for comparison
        
        article.save()
        return JsonResponse({'message': 'Article updated successfully!'}, status=200)
        

def about(request):
    return render(request, "about.html")
